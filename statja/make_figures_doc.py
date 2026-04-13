"""
Generate figures.docx — score screenshots grouped by cycle with minimal captions.
"""
import os, subprocess, time, re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _is_smooth(n):
    """Return True if n is of the form 2^a * 3^b (n >= 1)."""
    if n < 1:
        return False
    while n % 2 == 0:
        n //= 2
    while n % 3 == 0:
        n //= 3
    return n == 1


def add_caption_para(doc, text):
    """Add an italic caption paragraph; bold smooth occurrence numbers (≥ 8)
    found in the part after the last em-dash."""
    cap = doc.add_paragraph()
    cap.style = doc.styles['Normal']
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.space_after = Pt(10)

    # Split at last em-dash: prefix is plain italic, suffix gets smart bolding
    em = '\u2014'
    if em in text:
        idx = text.rfind(em)
        prefix = text[:idx + 1]
        suffix = text[idx + 1:]
    else:
        prefix, suffix = text, ''

    r = cap.add_run(prefix)
    r.italic = True
    r.font.size = Pt(9)

    for part in re.split(r'(\d+)', suffix):
        r = cap.add_run(part)
        r.italic = True
        r.font.size = Pt(9)
        if part.isdigit():
            n = int(part)
            if _is_smooth(n) and n >= 8:
                r.bold = True

    return cap

IMG_DIR = Path(__file__).parent / 'my_hand_png'
DST     = Path(__file__).parent / 'article.docx'

# ── catalogue: (filename_stem, caption) ──────────────────────────────────────
# caption = "Piece — pattern — N вхождений"
GROUPS = [
    ("Article Figures (§4)", [
        ("wtc1p07_krn_1_16_1_p1p1_inv",
            "Figure 1. BWV 852, WTC I Prelude 7 — 1/16; phase 1; +1+1 and inversion — 192 occurrences (2^6·3)"),
        ("wtc1f12_krn_1_16_1_p1p1p1",
            "Figure 2. BWV 857, WTC I Fugue 12 — 1/16; phase 1; +1+1+1 and inversion — 96 occurrences (2^5·3)"),
        ("bwv_944_2_xml_1_16_0_p2m1m1",
            "Figure 3. BWV 944, Fugue in A minor — 1/16; phase 0; +2−1−1 and inversion — 128 occurrences (2^7)"),
    ]),
    ("Inventions", [
        ("inven01_krn_1_16_1_m1m1m1p2m1p2m1",
            "BWV 772, Invention 1 — 1/16; −1−1−1+2−1+2−1 — 27 occurrences"),
        ("inven01_krn_1_16_1_m1p2m1",
            "BWV 772, Invention 1 — 1/16; −1+2−1 — 24 direct / 8 inverted / 32 total"),
        ("inven02_krn__1_4_1_16_0_m1m1m1_inv",
            "BWV 773, Invention 2 — (1/4)1/16; −1−1−1 with inversion — 36 occurrences (opening cell of the subject)"),
        ("inven02_krn_1_16_2_m1m1m1m1m1",
            "BWV 773, Invention 2 — 1/16; −1−1−1−1−1 — 16 occurrences"),
    ]),
    ("Sinfonias", [
        ("bwv787_xml_1_16_0_m1p1p1",
            "BWV 787, Sinfonia 1 — 1/16; −1+1+1 — 27 occurrences"),
        ("bwv787_xml_1_16_1_p1p1p1p1p1p1p1",
            "BWV 787, Sinfonia 1 — 1/16; +1+1+1+1+1+1+1 — 24 occurrences"),
        ("bwv788_xml_1_8_2_p0m2",
            "BWV 788, Sinfonia 2 — 1/8; +0−2 — 24 occurrences"),
        ("bwv788_xml_1_8_0_p1p1",
            "BWV 788, Sinfonia 2 — 1/8; +1+1 — 16 direct / 11 inverted / 27 total"),
        ("bwv788_xml_1_16_1_m1m1m1m1m1",
            "BWV 788, Sinfonia 2 — 1/16; −1−1−1−1−1 — 20 occurrences"),
    ]),
    ("Well-Tempered Clavier I", [
        ("wtc1f01_krn_1_8_1_p1p1p1",
            "BWV 846, Fugue 1 — 1/8; +1+1+1 — 24 occurrences (opening cell of the subject)"),
        ("wtc1f01_krn_1_32_2_m1m1",
            "BWV 846, Fugue 1 — 1/32; −1−1 — 24 direct / 3 inverted / 27 total"),
        ("wtc1f02_krn_1_8_1_m1m1",
            "BWV 847, Fugue 2 — 1/8; −1−1 — 18 occurrences"),
        ("wtc1f12_krn_1_16_1_p1p1p1",
            "BWV 857, Fugue 12 — 1/16; +1+1+1 — 96 occurrences"),
    ]),
    ("Well-Tempered Clavier II", [
        ("wtc2p01_krn_1_16_2_m1p1",
            "BWV 870, Prelude 1 — 1/16; −1+1 — 32 occurrences"),
        ("wtc2p01_krn_1_16_1_p1p1",
            "BWV 870, Prelude 1 — 1/16; +1+1 — 24 occurrences"),
        ("wtc2f01_krn_1_16_3_p1p1m2p1m2",
            "BWV 871, Fugue 1 — 1/16; +1+1−2+1−2 — 27 occurrences"),
        ("wtc2f01_krn_1_16_1_16__1_16_2_m1p1",
            "BWV 871, Fugue 1 — thematic seed 1/16,1/16,>1/16; −1+1 — 32 occurrences"),
        ("wtc2f11_krn_1_16_0_p1p1",
            "BWV 882, Fugue 11 — 1/16; +1+1 — 72 occurrences"),
        ("wtc2f11_krn_1_16_0_m1p1",
            "BWV 882, Fugue 11 — 1/16; −1+1 — 18 occurrences"),
        ("wtc2p22_krn__1_2_1_8_0_m1m1m1",
            "BWV 893, Prelude 22 — (1/2)1/8; −1−1−1 — 27 occurrences (opening cell of the subject)"),
        ("wtc2p22_krn_1_8_0_m1m1m1",
            "BWV 893, Prelude 22 — 1/8; −1−1−1 — 81 occurrences"),
        ("wtc2f22_krn_1_8_0_p1p1",
            "BWV 893, Fugue 22 — 1/8; +1+1 — 108 occurrences"),
        ("wtc2f22_krn_1_4_0_m4p1",
            "BWV 893, Fugue 22 — 1/4; −4+1 — 27 occurrences"),
        ("wtc2f19_krn_1_16_1_m1m1",
            "BWV 888, Fugue 19 — 1/16; −1−1 — 81 occurrences"),
        ("wtc2f19_krn_1_16_1_m1m1p1m3p1p1m1",
            "BWV 888, Fugue 19 — 1/16; −1−1+1−3+1+1−1 — 32 direct / 4 inverted / 36 total (second element of the subject)"),
    ]),
    ("French Suites", [
        ("bwv_812_Gigue6_xml_1_32_1_m1m1",
            "BWV 812, Gigue — 1/32; −1−1 — 25 direct / 16 inverted / 36 total"),
        ("bwv_812_Gigue6_xml_1_16_5_32_1_32_1_32_1_32_3_16_2_p1m1m1m1",
            "BWV 812, Gigue — thematic element +1−1−1−1−1+2 — 9 / 7 / 16"),
        ("french_suite_6_bach_french_suite_6_gigue_1_16_0_m1m1p1p1p1",
            "BWV 817, Gigue — 1/16; −1−1+1+1+1 — 18 occurrences"),
    ]),
    ("The Art of Fugue", [
        ("contrapunctusXI_xml_1_8_1_p1p0p0",
            "Contrapunctus XI — 1/8; +1+0+0 — 68 direct / 24 inverted / 81 total"),
        ("contrapunctusXI_xml_1_8_0_p0p0m2p1",
            "Contrapunctus XI — 1/8; +0+0−2+1 — 54 occurrences"),
        ("contrapunctusXI_xml_1_8_1_p1p0p0m2p1",
            "Contrapunctus XI — 1/8; +1+0+0−2+1 — 48 occurrences"),
        ("contrapunctusXI_xml_3_8_1_8_1_8_1_8_1_8_1_8_0_m2p1p0p0m2p1",
            "Contrapunctus XI — full motif −2+1+0+0−2+1 — 24 occurrences"),
        ("contrapunctusXI_xml__1_2_1_8_1_8__1_8_2_p1m1",
            "Contrapunctus XI — (1/2)1/8,1/8,>1/8; +1−1 with inversion — 24 / 48 occurrences"),
    ]),
    ("Organ Works", [
        ("vordeinenT_xml_1_8_0_p1p1",
            "BWV 668 «Vor deinen Thron» — 1/8; +1+1 — 81 occurrences"),
        ("bwv529_1_xml_1_16_0_m1m1",
            "BWV 529, movement I — 1/16; −1−1 — 128 occurrences"),
        ("bwv529_1_xml_1_16_1_m1p1p2m2m1p1",
            "BWV 529, movement I — thematic seed −1+1+2−2−1+1 — 36 occurrences"),
        ("bwv529_1_xml_1_8_1_m2p2",
            "BWV 529, movement I — 1/8; −2+2 — 36 occurrences"),
        ("bwv544_2_xml_1_8_1_p1p1p1",
            "BWV 544, movement II — 1/8; +1+1+1 — 96 occurrences"),
        ("bwv544_2_xml_1_8_0_m1p1p1",
            "BWV 544, movement II — 1/8; −1+1+1 — 81 occurrences"),
    ]),
    ("Violin Partitas", [
        ("partita1m4_krn_1_16_2_m1m1m1m1",
            "BWV 1002, Partita 1 — 1/16; −1−1−1−1 — 48 occurrences"),
        ("partita1m4_krn_1_16_1_p1p1p1p1p1p1p1",
            "BWV 1002, Partita 1 — 1/16; +1+1+1+1+1+1+1 — 32 occurrences"),
        ("partita1m4_krn_1_16_1_p1p1p1p1p1p1p1p1p1",
            "BWV 1002, Partita 1 — 1/16; +1+1+1+1+1+1+1+1+1 — 18 occurrences"),
    ]),
]

# ── build document ────────────────────────────────────────────────────────────

def make_doc():
    # open existing article.docx (generated by md_to_docx.py)
    doc = Document(str(DST))

    # page break before examples section
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    doc.add_page_break()

    # section heading
    t = doc.add_heading('Appendix: Annotated Score Excerpts', level=1)
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT

    intro = doc.add_paragraph(
        'Each excerpt shows the system containing the first occurrence of the motif '
        'and the system containing the last occurrence, separated by a row of three dots '
        '(\u00b7\u00b7\u00b7) to indicate omitted intervening material. '
        'Highlighted notes (orange) mark individual occurrences; '
        'bracketed boxes group the notes of each occurrence.'
    )
    intro.style = doc.styles['Normal']
    intro.paragraph_format.space_after = Pt(6)

    split_para = doc.add_paragraph(
        'Where a motif and its inversion are counted separately, '
        'the caption gives the format \u201cN direct / M inverted / K total\u201d. '
        'In several cases all three numbers are themselves smooth (of the form 2\u1d43\u00b73\u1d47): '
        'for instance, BWV\u00a0772 Invention\u00a01, pattern \u22121+2\u22121, yields '
        '24\u00a0direct / 8\u00a0inverted / 32\u00a0total '
        '(24\u00a0=\u00a02\u00b3\u00b73,\u00a08\u00a0=\u00a02\u00b3,\u00a032\u00a0=\u00a02\u2075). '
        'The smooth-number structure thus extends to the direct and inverted subcounts individually, '
        'not only to their union.'
    )
    split_para.style = doc.styles['Normal']
    split_para.paragraph_format.space_after = Pt(6)

    tool_para = doc.add_paragraph(
        'The excerpts were generated with kern_reader, an open-source score browser '
        'and motif-search tool developed for this study. It can be installed and run as follows:'
    )
    tool_para.style = doc.styles['Normal']
    tool_para.paragraph_format.space_after = Pt(4)

    for line in [
        'git clone https://github.com/vindomestic-oss/m_a',
        'cd m_a',
        'pip install verovio music21 pillow',
        'python kern_reader.py',
    ]:
        p = doc.add_paragraph(line)
        p.style = doc.styles['Normal']
        p.runs[0].font.name = 'Courier New'
        p.runs[0].font.size = Pt(10)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    for group_name, items in GROUPS:
        doc.add_heading(group_name, level=2)

        for stem, caption in items:
            img_path = IMG_DIR / (stem + '.png')
            if not img_path.exists():
                print(f'  MISSING: {img_path.name}')
                continue

            # all images same width = 16 cm
            doc.add_picture(str(img_path), width=Cm(16))
            # last paragraph is the picture; set caption below
            add_caption_para(doc, caption)

    doc.save(str(DST))
    print(f'Saved: {DST}')


if __name__ == '__main__':
    subprocess.run(['taskkill', '/F', '/IM', 'WINWORD.EXE'], capture_output=True)
    time.sleep(0.5)

    # step 1: generate article body
    import sys
    sys.path.insert(0, str(Path(__file__).parent))
    from md_to_docx import convert
    convert()

    # step 2: append figures section
    make_doc()

    os.startfile(str(DST))
