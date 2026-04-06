"""
Convert article.md → article.docx using python-docx.
Handles: headings H1-H3, paragraphs, bold/italic inline,
         horizontal rules (ignored), tables, block quotes.
"""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

SRC = "article.md"
DST = "article.docx"

IMG_DIR = "my_hand_png"

# Maps a substring of [FIGURE: ...] text → image filename in IMG_DIR
FIGURE_MAP = {
    "BWV 852":                         "wtc1p07_krn_1_16_1_p1p1_inv.png",
    "BWV 857":                         "wtc1f12_krn_1_16_1_p1p1p1.png",
    "BWV 944":                         "bwv_994_fugue_1_16-theme-seed.png",
}

# ── helpers ───────────────────────────────────────────────────────────────────

def set_run_fmt(run, bold=False, italic=False):
    run.bold   = bold   or None
    run.italic = italic or None


def add_inline(para, text):
    """Add text with inline **bold** and *italic* markers."""
    # Pattern: **bold**, *italic*, or plain
    pattern = re.compile(r'\*\*(.+?)\*\*|\*(.+?)\*|([^*]+)')
    for m in pattern.finditer(text):
        if m.group(1):
            r = para.add_run(m.group(1)); r.bold = True
        elif m.group(2):
            r = para.add_run(m.group(2)); r.italic = True
        else:
            para.add_run(m.group(3))


def add_table(doc, lines):
    """Parse a markdown table (lines starting with |) and add to doc."""
    rows = []
    for line in lines:
        if re.match(r'^\s*\|[-| :]+\|\s*$', line):
            continue  # separator row
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            add_inline(p, cell_text)
            if ri == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()  # spacing after table


# ── main ──────────────────────────────────────────────────────────────────────

def convert(src=SRC, dst=DST):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Default body font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Heading styles
    for level, pts in ((1, 16), (2, 14), (3, 12)):
        h = doc.styles[f'Heading {level}']
        h.font.name = 'Arial'
        h.font.size = Pt(pts)
        h.font.bold = True
        h.font.color.rgb = RGBColor(0, 0, 0)

    with open(src, encoding='utf-8') as f:
        raw = f.read()

    lines = raw.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # Horizontal rule → skip
        if re.match(r'^---+\s*$', line):
            i += 1
            continue

        # Heading
        m = re.match(r'^(#{1,3})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text  = m.group(2).strip()
            p = doc.add_heading('', level=level)
            p.clear()
            add_inline(p, text)
            i += 1
            continue

        # Table block: collect consecutive | lines
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, table_lines)
            continue

        # Blank line → paragraph break (skip)
        if not line.strip():
            i += 1
            continue

        # [FIGURE: ...] placeholder → insert image if available
        m = re.match(r'^\[FIGURE:\s*(.*)\]$', line.strip())
        if m:
            caption_text = m.group(1)
            img_file = None
            for key, fname in FIGURE_MAP.items():
                if key.lower() in caption_text.lower():
                    img_file = fname
                    break
            if img_file:
                import os
                img_path = os.path.join(os.path.dirname(src), IMG_DIR, img_file)
                if os.path.exists(img_path):
                    doc.add_picture(img_path, width=Cm(16))
                    cap = doc.add_paragraph(caption_text)
                    cap.style = doc.styles['Normal']
                    cap.runs[0].italic = True
                    cap.runs[0].font.size = Pt(9)
                    cap.paragraph_format.space_after = Pt(10)
                else:
                    p = doc.add_paragraph(f'[FIGURE — file not found: {img_file}]')
                    p.style = doc.styles['Normal']
            else:
                p = doc.add_paragraph(line)
                p.style = doc.styles['Normal']
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        add_inline(p, line)
        i += 1

    doc.save(dst)
    print(f"Saved: {dst}")


if __name__ == '__main__':
    import subprocess, time, os

    # Close Word if open
    subprocess.run(['taskkill', '/F', '/IM', 'WINWORD.EXE'],
                   capture_output=True)
    time.sleep(0.5)

    convert()

    # Reopen in Word
    os.startfile(os.path.abspath(DST))
